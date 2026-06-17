import os
from pypdf import PdfReader
from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LangchainDocument

def extract_text_from_pdf(file_path: str) -> list[LangchainDocument]:
    docs = []
    try:
        reader = PdfReader(file_path)
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                docs.append(LangchainDocument(
                    page_content=text,
                    metadata={"source": os.path.basename(file_path), "page": i + 1}
                ))
    except Exception as e:
        print(f"Error processing PDF {file_path}: {e}")
    return docs

def extract_text_from_docx(file_path: str) -> list[LangchainDocument]:
    docs = []
    try:
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        if text.strip():
            docs.append(LangchainDocument(
                page_content=text,
                metadata={"source": os.path.basename(file_path), "page": 1}
            ))
    except Exception as e:
        print(f"Error processing DOCX {file_path}: {e}")
    return docs

def extract_text_from_txt(file_path: str) -> list[LangchainDocument]:
    docs = []
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
            if text.strip():
                docs.append(LangchainDocument(
                    page_content=text,
                    metadata={"source": os.path.basename(file_path), "page": 1}
                ))
    except Exception as e:
        print(f"Error processing TXT {file_path}: {e}")
    return docs

def process_document(file_path: str, filename: str) -> list[LangchainDocument]:
    ext = os.path.splitext(filename)[1].lower()
    if ext == '.pdf':
        docs = extract_text_from_pdf(file_path)
    elif ext == '.docx':
        docs = extract_text_from_docx(file_path)
    elif ext == '.txt':
        docs = extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")
    
    # Chunking
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len
    )
    
    chunks = text_splitter.split_documents(docs)
    return chunks
